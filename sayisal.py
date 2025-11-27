# ==== SONUÇ TABLOLARI OLUŞTURMA (Spyder/Anaconda) ====
import os, pandas as pd
base = r"C:\Users\90530\Desktop\Akademi\7. Dönem\Sayısal Görüntü İşleme - Digital Image Processing\slaytlar\vize\BSDS500-master\BSDS500-master"
mm_path = os.path.join(base, "metrics_mean.csv")
m_path  = os.path.join(base, "metrics.csv")

# Okuma
mm = pd.read_csv(mm_path)
m  = pd.read_csv(m_path)

# Kolon adlarını normalize et (beklenen şemaya uydur)
def norm_cols(df):
    df = df.rename(columns={
        'color_norm':'RENK NORM.',
        'noise':'GÜRÜLTÜ',
        'denoise':'DENOISE YÖNTEMİ',
        'psnr':'PSNR (dB)',
        'ssim':'SSIM',
        'rmse':'RMSE',
        'mse':'MSE',
        'split':'VERİ BÖL.',
        'image':'GÖRSEL NO'
    })
    return df

mm = norm_cols(mm)
m  = norm_cols(m)

# Sıralama kuralı: SSIM↓, PSNR↓, RMSE↑, MSE↑
mm_sorted = mm.sort_values(by=['SSIM','PSNR (dB)','RMSE','MSE'],
                           ascending=[False, True, True, True])
m_sorted  = m.sort_values(by=['SSIM','PSNR (dB)','RMSE','MSE'],
                          ascending=[False, True, True, True])

# En iyi 5
mm_top5 = mm_sorted[['RENK NORM.','GÜRÜLTÜ','DENOISE YÖNTEMİ','PSNR (dB)','SSIM','RMSE','MSE']].head(5).copy()
m_top5  = m_sorted[['VERİ BÖL.','GÖRSEL NO','RENK NORM.','GÜRÜLTÜ','DENOISE YÖNTEMİ','PSNR (dB)','SSIM','RMSE','MSE']].head(5).copy()

# Sayısal biçim
for df in (mm_top5, m_top5):
    for c in ['PSNR (dB)','SSIM','RMSE','MSE']:
        df[c] = df[c].astype(float).round(3)

# Kaydet
mm_top5.to_csv(os.path.join(base, "TABLE_I_metrics_mean_TOP5.csv"), index=False)
m_top5.to_csv(os.path.join(base, "TABLE_II_metrics_TOP5.csv"), index=False)

print("TABLE I (metrics_mean, en iyi 5):\n", mm_top5.to_string(index=False))
print("\nTABLE II (metrics, en iyi 5):\n", m_top5.to_string(index=False))

# (Opsiyonel) Word çıktısı
try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    def to_docx(df, title, note, out_path):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        p_format = p.paragraph_format; p_format.space_after = Pt(6)

        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr = table.rows[0].cells
        for i,col in enumerate(df.columns):
            hdr[i].text = col.upper()

        for _,row in df.iterrows():
            cells = table.add_row().cells
            for i,val in enumerate(row):
                cells[i].text = str(val)

        p2 = doc.add_paragraph(note)
        p2.paragraph_format.space_before = Pt(4)
        doc.save(out_path)

    to_docx(mm_top5,
            "TABLE I. Yöntem Düzeyi Sonuçları (metrics_mean, en iyi 5)",
            "Not: Sıralama SSIM→PSNR→RMSE/MSE önceliğiyle; değerler ortalama performansı temsil eder.",
            os.path.join(base, "TABLE_I_metrics_mean_TOP5.docx"))
    to_docx(m_top5,
            "TABLE II. Görsel Düzeyi Örnek Sonuçlar (metrics, en iyi 5)",
            "Not: Her satır tek bir görüntü çıktısıdır; görsel no ve veri bölümü belirtilmiştir.",
            os.path.join(base, "TABLE_II_metrics_TOP5.docx"))
    print("\nDOCX dosyaları kaydedildi:", base)
except Exception as e:
    print("\n[Uyarı] Word çıktısı atlandı:", e, "\nCSV dosyaları hazır.")
